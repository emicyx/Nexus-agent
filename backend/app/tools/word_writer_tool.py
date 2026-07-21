"""Word 文档编写工具：生成 .docx 文件。"""
from typing import Any

from crewai.tools import BaseTool
from docx import Document
from pydantic import BaseModel, Field

from app.tools._file_utils import resolve_output_path


class WordWriterInput(BaseModel):
    filename: str = Field(
        ...,
        description="文件名（如 report.docx）",
    )
    title: str = Field(
        "",
        description="文档标题（可选）",
    )
    content: str = Field(
        ...,
        description="正文内容，用空行分隔段落",
    )
    sub_dir: str = Field(
        "",
        description="可选子目录，不填则写到 outputs 根目录",
    )


class WordWriterTool(BaseTool):
    """生成 Word .docx 文档。"""
    name: str = "write_word_doc"
    description: str = (
        "生成 Word .docx 文档。支持标题 + 多段落正文。"
        "触发时机：需要生成正式文档、报告、说明文件时使用。"
    )
    args_schema: type[BaseModel] = WordWriterInput

    def _run(
        self,
        filename: str = "",
        title: str = "",
        content: str = "",
        sub_dir: str = "",
        **kwargs: Any,
    ) -> str:
        if not filename:
            return "错误：filename 不能为空"
        if not filename.endswith(".docx"):
            filename += ".docx"

        doc = Document()
        if title:
            doc.add_heading(title, level=0)

        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        path = resolve_output_path(filename, sub_dir)
        doc.save(str(path))
        return f"Word 文档已生成: {path}"
